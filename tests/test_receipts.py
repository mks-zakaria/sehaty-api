"""Tests for the payment receipts.

Two things are worth testing on a document generator, and they are both about
what the paper claims.

The **amount in words** is the point of the exercise: on a Moroccan receipt it
is what makes the figure hard to alter afterwards, and French spells numbers
with enough exceptions — quatre-vingts, deux cents, mille, million — that a
naive speller writes something subtly wrong on every hundredth receipt. So the
speller is tested against the awkward cases, and against the totals actually
printed.

The **figures** are pinned to the commercial model. A receipt that disagrees
with the sales sheet is worse than no receipt: it is a contradiction in the
doctor's own file.

Skipped when the `print` extra is not installed:
    uv run --extra print pytest tests/test_receipts.py
"""

import importlib.util
import sys
from datetime import date
from pathlib import Path

import pytest

pytest.importorskip("reportlab", reason="install the 'print' extra")

_SCRIPTS = Path(__file__).resolve().parents[1] / "scripts"
sys.path.insert(0, str(_SCRIPTS))
_spec = importlib.util.spec_from_file_location("receipts", _SCRIPTS / "receipts.py")
receipts = importlib.util.module_from_spec(_spec)
sys.modules["receipts"] = receipts
_spec.loader.exec_module(receipts)

ISSUED = date(2026, 8, 1)

# Word stores column widths in twentieths of a point; 566.9 of them make a cm.
_W_ATTR = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w"


class TestAmountInWords:
    @pytest.mark.parametrize(
        ("amount", "expected"),
        [
            (1, "un dirham"),
            (2, "deux dirhams"),
            (17, "dix-sept dirhams"),
            (21, "vingt et un dirhams"),
            # The seventies and nineties are built on sixty and eighty.
            (71, "soixante et onze dirhams"),
            (77, "soixante-dix-sept dirhams"),
            (91, "quatre-vingt-onze dirhams"),
            # quatre-vingts keeps its s only when nothing follows.
            (80, "quatre-vingts dirhams"),
            (81, "quatre-vingt-un dirhams"),
            (80_000, "quatre-vingt mille dirhams"),
            # Same rule for cent.
            (200, "deux cents dirhams"),
            (201, "deux cent un dirhams"),
            (200_000, "deux cent mille dirhams"),
            (100, "cent dirhams"),
            # mille is invariable and never takes "un".
            (1_000, "mille dirhams"),
            (2_000, "deux mille dirhams"),
            # million is a noun: it pluralises, and a round one takes "de".
            (1_000_000, "un million de dirhams"),
            (2_000_000, "deux millions de dirhams"),
            (1_500_000, "un million cinq cent mille dirhams"),
        ],
    )
    def test_spells_the_awkward_cases(self, amount: int, expected: str) -> None:
        assert receipts.in_words(amount) == expected

    def test_spells_the_centimes_apart(self) -> None:
        assert receipts.in_words(1234.56) == (
            "mille deux cent trente-quatre dirhams et cinquante-six centimes"
        )
        assert receipts.in_words(2.01) == "deux dirhams et un centime"

    def test_rounds_to_the_centime_before_spelling(self) -> None:
        """The words and the figure must agree exactly.

        A float summing to 596.9999 prints as 597,00 — spelling out five hundred
        and ninety-six would put two different amounts on the same paper.
        """
        assert receipts.in_words(596.999) == receipts.in_words(597)

    def test_refuses_a_negative_amount(self) -> None:
        with pytest.raises(ValueError, match="negative"):
            receipts.in_words(-1)


class TestFormatting:
    def test_groups_thousands_the_french_way(self) -> None:
        assert receipts.money(2988) == "2 988,00"
        assert receipts.money(600) == "600,00"
        assert receipts.money(1234.5) == "1 234,50"


class TestTheThreeReceipts:
    def setup_method(self) -> None:
        self.presence, self.presence_rdv, self.renewal = receipts.examples(ISSUED, 1, 2026)

    def test_the_presence_pack_is_six_hundred_dirhams_once(self) -> None:
        assert self.presence.total == 600.0
        assert len(self.presence.items) == 1
        assert self.presence.items[0].quantity == 1
        assert "unique" in self.presence.period

    def test_presence_plus_rdv_bills_the_annual_rate_not_twelve_months(self) -> None:
        """600 + 1 990, inside the 1 200-3 500 first-year band.

        The annual rate is ten months, not twelve: two are offered. Charging
        12 × 199 would take 398 DH the sheet just promised to waive.
        """
        assert self.presence_rdv.total == 600.0 + 1990.0
        assert self.presence_rdv.total == 2590.0
        assert self.presence_rdv.total < 600.0 + 12 * 199.0
        assert 1200 <= self.presence_rdv.total <= 3500

    def test_the_renewal_is_the_quarterly_rate_from_the_sheet(self) -> None:
        assert self.renewal.total == 597.0
        assert self.renewal.total == 3 * 199.0
        assert "Virement" in self.renewal.method

    def test_the_words_match_the_figures(self) -> None:
        assert receipts.in_words(self.presence.total) == "six cents dirhams"
        assert receipts.in_words(self.presence_rdv.total) == (
            "deux mille cinq cent quatre-vingt-dix dirhams"
        )
        assert receipts.in_words(self.renewal.total) == ("cinq cent quatre-vingt-dix-sept dirhams")

    def test_receipt_numbers_are_sequential_within_the_year(self) -> None:
        numbers = [r.number for r in receipts.examples(ISSUED, 7, 2026)]
        assert numbers == ["SEH-2026-0007", "SEH-2026-0008", "SEH-2026-0009"]

    def test_every_example_is_marked_specimen(self) -> None:
        """A sample must never be passable as a real receipt."""
        assert all(r.specimen for r in receipts.examples(ISSUED, 1, 2026))

    def test_the_free_page_is_not_sold_twice(self) -> None:
        """The page is free forever; only the agenda is a subscription."""
        assert "gratuite" in " ".join(self.renewal.notes)


class TestAgreesWithTheSalesSheet:
    """The sheet is handed over first; the receipt must not requote it.

    This is the guard that was missing. The receipts originally billed a first
    year as 12 × 199, while the sheet offers the year at 1 990 — two months
    free. The doctor would have been charged 398 DH the sheet had just promised
    to waive, on the same visit.
    """

    def _sheet_text(self, tmp_path: Path) -> str:
        pytest.importorskip("pypdf")
        import sales_sheet
        from pypdf import PdfReader
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas as rl_canvas

        path = tmp_path / "sheet.pdf"
        pdf = rl_canvas.Canvas(str(path), pagesize=A4)
        sales_sheet.draw_recto(pdf)
        pdf.save()
        return PdfReader(path).pages[0].extract_text()

    def test_every_receipt_rate_is_quoted_on_the_sheet(self, tmp_path: Path) -> None:
        text = self._sheet_text(tmp_path).replace(" ", " ").replace("\xa0", " ")

        for amount in (
            receipts.PRESENCE_TTC,
            receipts.RDV_MONTHLY_TTC,
            receipts.RDV_QUARTER_TTC,
            receipts.RDV_YEAR_TTC,
        ):
            # The sheet writes whole dirhams: "600 DH", "1 990 DH".
            figure = receipts.money(amount).removesuffix(",00")
            assert figure in text, f"{figure} DH is not quoted on the sales sheet"


class TestWord:
    """The .docx exists so the signature and stamp can go on before printing."""

    def _document(self, tmp_path: Path, slug: str = "presence-rdv"):  # noqa: ANN202
        pytest.importorskip("docx")
        from docx import Document

        sys.argv = ["receipts.py", "--out", str(tmp_path), "--format", "docx"]
        assert receipts.main() == 0
        return Document(str(tmp_path / f"recu-{slug}.docx"))

    def test_writes_one_word_file_per_receipt(self, tmp_path: Path) -> None:
        self._document(tmp_path)
        written = sorted(p.name for p in tmp_path.glob("*.docx"))
        assert written == [
            "recu-presence-rdv.docx",
            "recu-presence.docx",
            "recu-rdv-renouvellement.docx",
        ]
        # --format docx means docx: no PDFs alongside it.
        assert list(tmp_path.glob("*.pdf")) == []

    def test_carries_the_same_figures_as_the_pdf(self, tmp_path: Path) -> None:
        document = self._document(tmp_path)
        text = "\n".join(p.text for p in document.paragraphs)

        assert "2 590,00 MAD" in text
        assert "deux mille cinq cent quatre-vingt-dix dirhams" in text
        assert "SPÉCIMEN" in text

    def test_holds_both_copies_with_a_cut_line(self, tmp_path: Path) -> None:
        document = self._document(tmp_path)
        text = "\n".join(p.text for p in document.paragraphs)

        assert "Exemplaire client" in text
        assert "Souche — Sehaty" in text
        assert "découper ici" in text

    def test_leaves_an_empty_box_to_drop_a_signature_into(self, tmp_path: Path) -> None:
        """A cell, not a floating image anchor: dropping a picture into a cell
        keeps it where it was put instead of sliding over the text."""
        document = self._document(tmp_path)
        signatures = document.tables[-1]

        assert [cell.text for cell in signatures.rows[0].cells] == ["", ""]
        assert [cell.text for cell in signatures.rows[1].cells] == [
            "Le client",
            "Pour Sehaty — cachet et signature",
        ]

    def test_the_columns_are_pinned_so_word_cannot_reflow_them(self, tmp_path: Path) -> None:
        """Autofit collapses the désignation column and pushes the souche onto a
        second page, which breaks the cut-in-half layout."""
        document = self._document(tmp_path)
        items = document.tables[0]

        assert items.autofit is False
        widths = [round(cell.width.cm, 1) for cell in items.rows[0].cells]
        assert widths == [11.3, 1.3, 2.6, 2.6]
        # The grid has to agree, or a fixed layout still renders the old widths.
        grid = [round(int(col.get(_W_ATTR)) / 566.9, 1) for col in items._tbl.tblGrid.gridCol_lst]
        assert grid == widths


class TestRendering:
    def test_writes_one_pdf_per_receipt(self, tmp_path: Path) -> None:
        sys.argv = ["receipts.py", "--out", str(tmp_path), "--date", "2026-08-01"]
        assert receipts.main() == 0

        written = sorted(p.name for p in tmp_path.glob("*.pdf"))
        assert written == [
            "recu-presence-rdv.pdf",
            "recu-presence.pdf",
            "recu-rdv-renouvellement.pdf",
        ]
        for pdf in tmp_path.glob("*.pdf"):
            assert pdf.stat().st_size > 1000, pdf

    def test_each_page_carries_both_copies(self, tmp_path: Path) -> None:
        """Client copy and souche on one sheet, so the numbers cannot drift."""
        pytest.importorskip("pypdf")
        from pypdf import PdfReader

        sys.argv = ["receipts.py", "--out", str(tmp_path)]
        receipts.main()

        text = PdfReader(tmp_path / "recu-presence.pdf").pages[0].extract_text()
        assert "EXEMPLAIRE CLIENT" in text.upper()
        assert "SOUCHE" in text.upper()
        assert text.upper().count("REÇU DE PAIEMENT") == 2

    def test_the_issuing_company_is_named_on_the_paper(self, tmp_path: Path) -> None:
        """A commercial document a doctor keeps has to say who billed them."""
        pytest.importorskip("pypdf")
        from pypdf import PdfReader

        sys.argv = ["receipts.py", "--out", str(tmp_path)]
        receipts.main()

        text = PdfReader(tmp_path / "recu-presence.pdf").pages[0].extract_text()
        assert receipts.COMPANY["name"] in text
        assert "ICE" in text
