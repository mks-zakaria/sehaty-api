#!/usr/bin/env python
"""Generate payment receipts (reçus) for the packs, A4, French.

The paper a doctor keeps after handing over cash. It is the only document from
this business they will ever file, so it has to look like a receipt and not like
a printout: an issuer with real registration numbers, a number, a date, what was
paid for, the amount in figures **and in words**, and two signatures.

Printed in two parts on one page. The upper half is the doctor's; the lower half
is the souche you keep, identical down to the number. Cut along the dashes. That
is how paper receipts work here, and it is also the only way the receipt number
on your copy and theirs cannot drift apart.

**The amount in words is not decoration.** On a Moroccan receipt it is what
makes the figure hard to alter after the fact, and its absence is the first
thing anyone contesting a payment will point at. It is spelled out in full
French, so the speller is tested.

**Every figure is TTC**, matching the sales sheet and the invoice. Doctors
cannot recover TVA on a medical practice, so a price that grows 20% between the
desk and the paper is a trust problem on the first transaction.

Receipt numbers are sequential per year (SEH-2026-0001). Pass `--number` to set
the first one; the examples below are numbered from 1 and are marked SPECIMEN so
that a sample can never be mistaken for a real receipt.

Signature blocks are left blank on purpose. Print the PDF and sign it by hand,
or open the .docx, click inside the signature cell, drop a scanned signature in
and print that. Both formats carry the same figures from the same constants, so
neither can quietly drift from the other.

Usage:
    cd sehaty-api
    uv run --extra print python scripts/receipts.py --out ./print
    uv run --extra print python scripts/receipts.py --out ./print --format docx

French only, by design: this is the language of Moroccan professional paperwork.
(reportlab also cannot shape Arabic — see `print_assets.py`.)
"""

from __future__ import annotations

import argparse
import sys
from dataclasses import dataclass, field
from datetime import date
from pathlib import Path

from reportlab.lib.colors import Color, HexColor
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.pdfgen import canvas

# The issuing identity lives in one place. Two documents disagreeing about the
# ICE is exactly the kind of thing that gets noticed after 200 are printed.
sys.path.insert(0, str(Path(__file__).resolve().parent))
from sales_sheet import (  # noqa: E402
    BRAND,
    BRAND_DARK,
    COMPANY,
    CONTACT,
    HAIRLINE,
    MUTED,
    wrap,
)

INK = HexColor("#0f172a")
PAPER = HexColor("#f8fafc")

CURRENCY = "MAD"

# The commercial model, in one place. These must stay in step with sales_sheet:
# the sheet is handed over first, and a receipt that recomputes a price the
# doctor was just quoted contradicts you in their own file.
#
# The agenda is never billed monthly — collecting from twenty cabinets every
# month does not scale — so only the quarter and the year have prices. The year
# is not twelve months' worth: it is ten, with two offered.
PRESENCE_TTC = 600.0
RDV_MONTHLY_TTC = 199.0
RDV_QUARTER_TTC = 597.0
RDV_YEAR_TTC = 1990.0
RDV_MONTHS_QUARTER = 3

# Mentions that belong on a receipt for a cash payment.
TAX_NOTE = "Montant TTC (TVA 20 % incluse). Prix nets à payer."
CASH_NOTE = (
    "Reçu vaut quittance pour la période indiquée. En cas de règlement par "
    "chèque, la quittance n'est définitive qu'après encaissement."
)


@dataclass(frozen=True)
class LineItem:
    """One thing paid for."""

    label: str
    detail: str
    quantity: int
    unit_price: float

    @property
    def total(self) -> float:
        return self.quantity * self.unit_price


@dataclass(frozen=True)
class Receipt:
    """Everything printed on one receipt."""

    number: str
    issued_on: date
    # Who paid. Free text: the receipt is written for whoever is at the desk,
    # and plenty of cabinets pay under the practice name rather than the doctor's.
    payer: str
    payer_detail: str
    subject: str
    items: list[LineItem]
    method: str
    period: str | None = None
    specimen: bool = False
    notes: list[str] = field(default_factory=list)

    @property
    def total(self) -> float:
        return sum(item.total for item in self.items)


# --------------------------------------------------------------------------
# Amount in words
# --------------------------------------------------------------------------

_UNITS = [
    "zéro",
    "un",
    "deux",
    "trois",
    "quatre",
    "cinq",
    "six",
    "sept",
    "huit",
    "neuf",
    "dix",
    "onze",
    "douze",
    "treize",
    "quatorze",
    "quinze",
    "seize",
]
_TENS = {
    2: "vingt",
    3: "trente",
    4: "quarante",
    5: "cinquante",
    6: "soixante",
    8: "quatre-vingt",
}


def _under_hundred(n: int, *, final: bool = True) -> str:
    """0-99 in French, including the seventies and eighties.

    ``final`` is False when another word follows the number, which is what
    decides the plural: *quatre-vingts* alone, *quatre-vingt mille*.
    """
    if n < 17:
        return _UNITS[n]
    if n < 20:
        return f"dix-{_UNITS[n - 10]}"

    tens, unit = divmod(n, 10)
    # 70-79 and 90-99 are built on sixty and eighty: soixante-douze, quatre-vingt-douze.
    if tens in (7, 9):
        base = _TENS[tens - 1]
        rest = _under_hundred(10 + unit)
        # 71 alone keeps the conjunction: soixante et onze.
        joiner = " et " if unit == 1 and tens == 7 else "-"
        return f"{base}{joiner}{rest}"

    word = _TENS[tens]
    if unit == 0:
        return f"{word}s" if tens == 8 and final else word
    if unit == 1 and tens != 8:
        return f"{word} et un"
    return f"{word}-{_UNITS[unit]}"


def _under_thousand(n: int, *, final: bool = True) -> str:
    hundreds, rest = divmod(n, 100)
    if hundreds == 0:
        return _under_hundred(rest, final=final)

    # "cent" not "un cent"; plural only when it ends the number: deux cents,
    # deux cent trois, and deux cent mille.
    head = "cent" if hundreds == 1 else f"{_UNITS[hundreds]} cent"
    if rest == 0:
        return head if hundreds == 1 or not final else f"{head}s"
    return f"{head} {_under_hundred(rest)}"


def in_words(amount: float) -> str:
    """A dirham amount spelled out, the way a receipt has to carry it.

    Rounded to the centime first: the words and the figure must agree exactly,
    and a float that prints as 597,00 while summing to 596,999 would spell out
    the wrong number.
    """
    if amount < 0:
        raise ValueError("a receipt cannot be for a negative amount")

    centimes_total = int(round(amount * 100))
    dirhams, centimes = divmod(centimes_total, 100)

    if dirhams >= 1_000_000_000:
        raise ValueError("amount out of range for the speller")

    parts: list[str] = []
    millions, rest = divmod(dirhams, 1_000_000)
    if millions:
        # "un million" keeps its article, unlike mille. Cent and vingt do take
        # the plural here — million is a noun, not a numeral: deux cents millions.
        parts.append(f"{_under_thousand(millions)} million" + ("s" if millions > 1 else ""))

    thousands, units = divmod(rest, 1000)
    if thousands:
        # mille is invariable and takes no "un": mille, deux mille. Nothing
        # before it pluralises either: quatre-vingt mille, deux cent mille.
        parts.append(
            "mille" if thousands == 1 else f"{_under_thousand(thousands, final=False)} mille"
        )
    if units or not parts:
        parts.append(_under_thousand(units))

    words = " ".join(parts)
    # "million" is a noun, so a round million is followed by "de": un million de
    # dirhams, but un million cinq cent mille dirhams.
    if millions and rest == 0:
        spelled = f"{words} de dirhams"
    else:
        label = "dirham" if dirhams <= 1 else "dirhams"
        spelled = f"{words} {label}"
    if centimes:
        spelled += f" et {_under_hundred(centimes)} centime" + ("s" if centimes > 1 else "")
    return spelled


def money(amount: float) -> str:
    """1234.5 -> '1 234,50' — French grouping, as printed on the receipt."""
    formatted = f"{amount:,.2f}".replace(",", " ").replace(".", ",")
    return formatted


# --------------------------------------------------------------------------
# The three examples
# --------------------------------------------------------------------------


def examples(issued_on: date, first_number: int, year: int) -> list[Receipt]:
    """The three receipts you will actually write, in the order you meet them."""

    def number(offset: int) -> str:
        return f"SEH-{year}-{first_number + offset:04d}"

    presence = Receipt(
        number=number(0),
        issued_on=issued_on,
        payer="Dr Amina Bennani",
        payer_detail="Dentiste — Casablanca, Maârif",
        subject="Pack Présence en ligne",
        items=[
            LineItem(
                label="Pack Présence en ligne",
                detail=(
                    "Page professionnelle sur sehaty.ma, plaque QR pour la salle "
                    "d'attente, 100 cartes de poche, fiche Google corrigée, "
                    "photos du cabinet, référencement dans l'annuaire 12 mois, "
                    "statistiques mensuelles."
                ),
                quantity=1,
                unit_price=PRESENCE_TTC,
            ),
        ],
        method="Espèces",
        period="Mise en service — paiement unique",
        specimen=True,
        notes=[
            "La page reste en ligne sans abonnement. Aucun prélèvement "
            "automatique n'est mis en place.",
        ],
    )

    presence_rdv = Receipt(
        number=number(1),
        issued_on=issued_on,
        payer="Dr Karim Alami",
        payer_detail="Cardiologue — Rabat, Agdal",
        subject="Pack Présence en ligne + système de rendez-vous",
        items=[
            LineItem(
                label="Pack Présence en ligne",
                detail="Mise en service, paiement unique.",
                quantity=1,
                unit_price=PRESENCE_TTC,
            ),
            LineItem(
                label="Système de rendez-vous en ligne — abonnement annuel",
                detail=(
                    "Agenda en ligne, confirmations et rappels automatiques aux "
                    "patients. Année réglée d'avance : 12 mois, dont 2 offerts "
                    "(199 DH TTC/mois, tarif fondateur)."
                ),
                quantity=1,
                unit_price=RDV_YEAR_TTC,
            ),
        ],
        method="Espèces",
        period=(
            f"Abonnement du {issued_on:%d/%m/%Y} "
            f"au {issued_on.replace(year=issued_on.year + 1):%d/%m/%Y}"
        ),
        specimen=True,
        notes=[
            "Tarif fondateur bloqué 24 mois. Sans reconduction tacite : "
            "l'abonnement s'arrête si rien n'est réglé à l'échéance.",
        ],
    )

    renewal = Receipt(
        number=number(2),
        issued_on=issued_on,
        payer="Dr Karim Alami",
        payer_detail="Cardiologue — Rabat, Agdal",
        subject="Abonnement rendez-vous — renouvellement trimestriel",
        items=[
            LineItem(
                label="Système de rendez-vous en ligne — trimestre",
                detail=(
                    f"Renouvellement, {RDV_MONTHS_QUARTER} mois réglés d'avance "
                    "(199 DH TTC/mois, tarif fondateur)."
                ),
                quantity=1,
                unit_price=RDV_QUARTER_TTC,
            ),
        ],
        method="Virement bancaire",
        period="Trimestre du 01/10/2026 au 31/12/2026",
        specimen=True,
        notes=[
            "La page professionnelle reste gratuite et n'est pas concernée par cet abonnement.",
        ],
    )

    return [presence, presence_rdv, renewal]


# --------------------------------------------------------------------------
# Drawing
# --------------------------------------------------------------------------


def _label(pdf: canvas.Canvas, text: str, x: float, y: float) -> None:
    pdf.setFillColor(MUTED)
    pdf.setFont("Helvetica", 7)
    pdf.drawString(x, y, text.upper())


def draw_receipt(
    pdf: canvas.Canvas,
    receipt: Receipt,
    *,
    origin_y: float,
    copy_label: str,
    width: float,
) -> None:
    """Draw one copy of the receipt, its top edge at ``origin_y``."""
    left = 18 * mm
    right = width - 18 * mm
    inner = right - left
    y = origin_y

    # Brand bar, so the two halves read as separate documents once cut.
    pdf.setFillColor(BRAND)
    pdf.rect(left, y - 1.5 * mm, inner, 1.5 * mm, stroke=0, fill=1)
    y -= 9 * mm

    pdf.setFillColor(BRAND_DARK)
    pdf.setFont("Helvetica-Bold", 15)
    pdf.drawString(left, y, "REÇU DE PAIEMENT")

    pdf.setFillColor(MUTED)
    pdf.setFont("Helvetica", 8)
    pdf.drawRightString(right, y + 4 * mm, copy_label.upper())
    pdf.setFillColor(INK)
    pdf.setFont("Helvetica-Bold", 11)
    pdf.drawRightString(right, y - 1 * mm, f"N° {receipt.number}")
    pdf.setFillColor(MUTED)
    pdf.setFont("Helvetica", 9)
    pdf.drawRightString(right, y - 6 * mm, f"Casablanca, le {receipt.issued_on:%d/%m/%Y}")

    # Issuer, small: the doctor knows who you are, the tax office does not.
    y -= 6 * mm
    pdf.setFillColor(MUTED)
    pdf.setFont("Helvetica", 7.5)
    pdf.drawString(left, y, f"{COMPANY['name']} — {COMPANY['address']}")
    y -= 3.6 * mm
    pdf.drawString(
        left,
        y,
        f"RC {COMPANY['rc']} — ICE {COMPANY['ice']} — IF {COMPANY['if']} — {CONTACT}",
    )

    y -= 6.5 * mm
    _label(pdf, "Reçu de", left, y)
    y -= 5 * mm
    pdf.setFillColor(INK)
    pdf.setFont("Helvetica-Bold", 11)
    pdf.drawString(left, y, receipt.payer)
    pdf.setFillColor(MUTED)
    pdf.setFont("Helvetica", 9)
    pdf.drawString(
        left + pdf.stringWidth(receipt.payer, "Helvetica-Bold", 11) + 3 * mm,
        y,
        receipt.payer_detail,
    )

    y -= 6 * mm
    _label(pdf, "Au titre de", left, y)
    y -= 5 * mm
    pdf.setFillColor(INK)
    pdf.setFont("Helvetica", 10.5)
    pdf.drawString(left, y, receipt.subject)

    # Items.
    y -= 6.5 * mm
    pdf.setStrokeColor(HAIRLINE)
    pdf.setLineWidth(0.6)
    pdf.line(left, y, right, y)
    y -= 4.5 * mm
    _label(pdf, "Désignation", left, y)
    _label(pdf, "Qté", right - 52 * mm, y)
    _label(pdf, "P.U. TTC", right - 38 * mm, y)
    _label(pdf, "Total TTC", right - 18 * mm, y)
    y -= 2 * mm
    pdf.line(left, y, right, y)

    for item in receipt.items:
        y -= 5.5 * mm
        pdf.setFillColor(INK)
        pdf.setFont("Helvetica-Bold", 9.5)
        pdf.drawString(left, y, item.label)
        pdf.setFont("Helvetica", 9.5)
        pdf.drawString(right - 52 * mm, y, str(item.quantity))
        pdf.drawString(right - 38 * mm, y, money(item.unit_price))
        pdf.drawRightString(right, y, money(item.total))

        pdf.setFillColor(MUTED)
        pdf.setFont("Helvetica", 8)
        for line in wrap(pdf, item.detail, "Helvetica", 8, inner - 58 * mm):
            y -= 3.4 * mm
            pdf.drawString(left, y, line)

    y -= 4 * mm
    pdf.setStrokeColor(HAIRLINE)
    pdf.line(left, y, right, y)

    # Total, boxed — the number the eye goes to first.
    y -= 10 * mm
    pdf.setFillColor(PAPER)
    pdf.setStrokeColor(BRAND)
    pdf.setLineWidth(0.8)
    pdf.roundRect(right - 70 * mm, y, 70 * mm, 11 * mm, 1.5 * mm, stroke=1, fill=1)
    pdf.setFillColor(BRAND_DARK)
    pdf.setFont("Helvetica-Bold", 9)
    pdf.drawString(right - 66 * mm, y + 4 * mm, "TOTAL RÉGLÉ")
    pdf.setFont("Helvetica-Bold", 13)
    pdf.drawRightString(right - 4 * mm, y + 3.4 * mm, f"{money(receipt.total)} {CURRENCY}")

    pdf.setFillColor(MUTED)
    pdf.setFont("Helvetica", 8)
    pdf.drawString(left, y + 7 * mm, f"Mode de règlement : {receipt.method}")
    if receipt.period:
        pdf.drawString(left, y + 3 * mm, receipt.period)

    # Amount in words: what makes the figure hard to alter afterwards.
    y -= 5 * mm
    pdf.setFillColor(INK)
    pdf.setFont("Helvetica-Oblique", 8.5)
    spelled = f"Arrêté le présent reçu à la somme de : {in_words(receipt.total)}."
    for line in wrap(pdf, spelled, "Helvetica-Oblique", 8.5, inner):
        pdf.drawString(left, y, line)
        y -= 4 * mm

    pdf.setFillColor(MUTED)
    pdf.setFont("Helvetica", 7)
    for note in [*receipt.notes, TAX_NOTE, CASH_NOTE]:
        for line in wrap(pdf, note, "Helvetica", 7, inner):
            y -= 3.2 * mm
            pdf.drawString(left, y, line)

    # Signatures — left blank, signed and stamped by hand. The boxes hang below
    # the current baseline, so the gap above has to clear their full height or
    # they are drawn straight through the notes.
    y -= 6 * mm
    box_h = 13 * mm
    box_w = 62 * mm
    for x, caption in ((left, "Le client"), (right - box_w, "Pour Sehaty — cachet et signature")):
        pdf.setStrokeColor(HAIRLINE)
        pdf.setLineWidth(0.6)
        pdf.rect(x, y - box_h, box_w, box_h, stroke=1, fill=0)
        pdf.setFillColor(MUTED)
        pdf.setFont("Helvetica", 7)
        pdf.drawString(x, y - box_h - 3.5 * mm, caption)

    if receipt.specimen:
        # Across the middle of this copy, clear of the amount in words — a
        # specimen has to be unusable, not unreadable.
        _watermark(pdf, left + inner / 2, (origin_y + y) / 2)


def _watermark(pdf: canvas.Canvas, x: float, y: float) -> None:
    """SPECIMEN across the copy, so a sample cannot be passed off as a receipt."""
    pdf.saveState()
    pdf.setFillColor(Color(0.85, 0.15, 0.15, alpha=0.10))
    pdf.setFont("Helvetica-Bold", 42)
    pdf.translate(x, y)
    pdf.rotate(18)
    pdf.drawCentredString(0, 0, "SPÉCIMEN")
    pdf.restoreState()


def draw_page(pdf: canvas.Canvas, receipt: Receipt) -> None:
    """Client copy above, souche below, cut line between."""
    width, height = A4

    draw_receipt(
        pdf, receipt, origin_y=height - 12 * mm, copy_label="Exemplaire client", width=width
    )

    middle = height / 2
    pdf.setStrokeColor(HAIRLINE)
    pdf.setLineWidth(0.7)
    pdf.setDash(2, 3)
    pdf.line(12 * mm, middle, width - 12 * mm, middle)
    pdf.setDash()
    pdf.setFillColor(MUTED)
    pdf.setFont("Helvetica", 6.5)
    pdf.drawCentredString(width / 2, middle + 1.6 * mm, "— — —  découper ici  — — —")

    draw_receipt(pdf, receipt, origin_y=middle - 6 * mm, copy_label="Souche — Sehaty", width=width)

    pdf.showPage()


# --------------------------------------------------------------------------
# Word
# --------------------------------------------------------------------------
#
# The same receipt as an editable .docx, because the signature and the stamp go
# on after the fact — either scanned in, or typed over before printing. The PDF
# is the thing you print; the Word file is the thing you finish.
#
# Laid out with tables rather than tabs. Word reflows text as soon as anyone
# edits it, and a tab-aligned column silently falls apart the first time a
# doctor's name runs long. A table cell does not move.


def _docx_shade(cell, hex_colour: str) -> None:  # noqa: ANN001
    """Fill a table cell. python-docx has no API for it, so write the XML."""
    from docx.oxml.ns import qn
    from docx.oxml.parser import OxmlElement

    shade = OxmlElement("w:shd")
    shade.set(qn("w:val"), "clear")
    shade.set(qn("w:fill"), hex_colour)
    cell._tc.get_or_add_tcPr().append(shade)


def _docx_fix_widths(table, widths) -> None:  # noqa: ANN001
    """Pin the column widths.

    Word autofits by default, which collapses the désignation column to fit the
    numbers and then wraps the description over six lines — enough to push the
    second copy onto a second page. Fixing the layout needs both the table-level
    ``tblLayout`` element and a width on every cell; setting one without the
    other is silently ignored.
    """
    from docx.oxml.ns import qn

    # `autofit = False` writes the tblLayout element itself, in the position the
    # schema requires. Appending one by hand puts it after tblLook, out of
    # order, and both Word and LibreOffice then ignore it.
    table.autofit = False

    for row in table.rows:
        for cell, width in zip(row.cells, widths, strict=True):
            cell.width = width

    # Under a fixed layout the renderer lays the columns out from tblGrid, which
    # still holds the equal widths from when the table was created. Setting only
    # the cell widths changes nothing on the page.
    for column, width in zip(table._tbl.tblGrid.gridCol_lst, widths, strict=True):
        column.set(qn("w:w"), str(width.twips))


def _docx_run(paragraph, text: str, *, size: int, bold: bool = False, colour=None, italic=False):  # noqa: ANN001, ANN202
    from docx.shared import Pt

    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if colour is not None:
        run.font.color.rgb = colour
    return run


def _docx_copy(document, receipt: Receipt, *, copy_label: str) -> None:  # noqa: ANN001
    """One copy of the receipt into an open document."""
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor

    ink = RGBColor(0x0F, 0x17, 0x2A)
    brand = RGBColor(0x1B, 0x3D, 0x5E)
    muted = RGBColor(0x64, 0x74, 0x8B)

    head = document.add_paragraph()
    head.paragraph_format.space_after = Pt(1)
    _docx_run(head, "REÇU DE PAIEMENT", size=14, bold=True, colour=brand)

    meta = document.add_paragraph()
    meta.paragraph_format.space_after = Pt(1)
    _docx_run(meta, f"{copy_label}  ·  ", size=8, colour=muted)
    _docx_run(meta, f"N° {receipt.number}", size=10, bold=True, colour=ink)
    _docx_run(meta, f"  ·  Casablanca, le {receipt.issued_on:%d/%m/%Y}", size=9, colour=muted)

    issuer = document.add_paragraph()
    issuer.paragraph_format.space_after = Pt(5)
    _docx_run(issuer, f"{COMPANY['name']} — {COMPANY['address']}\n", size=7.5, colour=muted)
    _docx_run(
        issuer,
        f"RC {COMPANY['rc']} — ICE {COMPANY['ice']} — IF {COMPANY['if']} — {CONTACT}",
        size=7.5,
        colour=muted,
    )

    for label, value, detail in (
        ("REÇU DE", receipt.payer, receipt.payer_detail),
        ("AU TITRE DE", receipt.subject, ""),
    ):
        block = document.add_paragraph()
        block.paragraph_format.space_after = Pt(3)
        _docx_run(block, f"{label}\n", size=7, colour=muted)
        _docx_run(block, value, size=11, bold=True, colour=ink)
        if detail:
            _docx_run(block, f"   {detail}", size=9, colour=muted)

    # Items.
    item_widths = (Cm(11.3), Cm(1.3), Cm(2.6), Cm(2.6))
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for cell, text in zip(
        table.rows[0].cells,
        ("DÉSIGNATION", "QTÉ", "P.U. TTC", "TOTAL TTC"),
        strict=True,
    ):
        _docx_shade(cell, "F1F5F9")
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        _docx_run(paragraph, text, size=7, bold=True, colour=muted)

    for item in receipt.items:
        cells = table.add_row().cells
        first = cells[0].paragraphs[0]
        first.paragraph_format.space_after = Pt(0)
        _docx_run(first, item.label, size=9.5, bold=True, colour=ink)
        detail = cells[0].add_paragraph()
        detail.paragraph_format.space_before = Pt(0)
        _docx_run(detail, item.detail, size=7.5, colour=muted)
        for cell, text in zip(
            cells[1:],
            (str(item.quantity), money(item.unit_price), money(item.total)),
            strict=True,
        ):
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            _docx_run(paragraph, text, size=9.5, colour=ink)

    _docx_fix_widths(table, item_widths)

    total_row = document.add_paragraph()
    total_row.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    total_row.paragraph_format.space_before = Pt(4)
    total_row.paragraph_format.space_after = Pt(2)
    _docx_run(total_row, "TOTAL RÉGLÉ   ", size=9, bold=True, colour=muted)
    _docx_run(total_row, f"{money(receipt.total)} {CURRENCY}", size=14, bold=True, colour=brand)

    method = document.add_paragraph()
    method.paragraph_format.space_after = Pt(3)
    _docx_run(method, f"Mode de règlement : {receipt.method}", size=8, colour=muted)
    if receipt.period:
        _docx_run(method, f"\n{receipt.period}", size=8, colour=muted)

    spelled = document.add_paragraph()
    spelled.paragraph_format.space_after = Pt(4)
    _docx_run(
        spelled,
        f"Arrêté le présent reçu à la somme de : {in_words(receipt.total)}.",
        size=8.5,
        italic=True,
        colour=ink,
    )

    for note in [*receipt.notes, TAX_NOTE, CASH_NOTE]:
        line = document.add_paragraph()
        line.paragraph_format.space_after = Pt(1)
        _docx_run(line, note, size=6.5, colour=muted)

    # Signatures. Two empty cells sized for a scanned signature: click inside,
    # Insertion > Images, and it lands in the box instead of floating over the
    # text the way a free-floating picture does.
    signatures = document.add_table(rows=2, cols=2)
    signatures.style = "Table Grid"
    signatures.rows[0].height = Cm(1.5)
    for cell in signatures.rows[0].cells:
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    for cell, caption in zip(
        signatures.rows[1].cells,
        ("Le client", "Pour Sehaty — cachet et signature"),
        strict=True,
    ):
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        _docx_run(paragraph, caption, size=7, colour=muted)
    _docx_fix_widths(signatures, (Cm(8.9), Cm(8.9)))

    if receipt.specimen:
        mark = document.add_paragraph()
        mark.alignment = WD_ALIGN_PARAGRAPH.CENTER
        mark.paragraph_format.space_before = Pt(2)
        _docx_run(mark, "SPÉCIMEN", size=10, bold=True, colour=RGBColor(0xD9, 0x26, 0x26))


def write_docx(receipt: Receipt, path: Path) -> None:
    """The receipt as an editable Word file, both copies, cut line between."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor

    document = Document()
    document.core_properties.title = f"Sehaty — Reçu {receipt.number}"

    section = document.sections[0]
    section.top_margin = section.bottom_margin = Cm(1.0)
    section.left_margin = section.right_margin = Cm(1.6)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(2)

    _docx_copy(document, receipt, copy_label="Exemplaire client")

    cut = document.add_paragraph()
    cut.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cut.paragraph_format.space_before = Pt(6)
    cut.paragraph_format.space_after = Pt(6)
    _docx_run(
        cut,
        "— — — — — — — — — —  découper ici  — — — — — — — — — —",
        size=7,
        colour=RGBColor(0x64, 0x74, 0x8B),
    )

    _docx_copy(document, receipt, copy_label="Souche — Sehaty")
    document.save(str(path))


SLUGS = ("presence", "presence-rdv", "rdv-renouvellement")
FORMATS = ("pdf", "docx", "both")


def main() -> int:
    parser = argparse.ArgumentParser(description="Generate example payment receipts.")
    parser.add_argument("--out", type=Path, default=Path("print"), help="Output directory")
    parser.add_argument(
        "--number",
        type=int,
        default=1,
        help="First receipt number in the year's sequence (default 1)",
    )
    parser.add_argument(
        "--date",
        type=date.fromisoformat,
        default=date.today(),
        help="Issue date, YYYY-MM-DD (default today)",
    )
    parser.add_argument(
        "--format",
        choices=FORMATS,
        default="both",
        help="pdf to print as-is, docx to sign and edit first (default both)",
    )
    args = parser.parse_args()

    args.out.mkdir(parents=True, exist_ok=True)
    receipts = examples(args.date, args.number, args.date.year)

    for slug, receipt in zip(SLUGS, receipts, strict=True):
        written: list[Path] = []
        if args.format in ("pdf", "both"):
            path = args.out / f"recu-{slug}.pdf"
            pdf = canvas.Canvas(str(path), pagesize=A4)
            pdf.setTitle(f"Sehaty — Reçu {receipt.number}")
            draw_page(pdf, receipt)
            pdf.save()
            written.append(path)
        if args.format in ("docx", "both"):
            path = args.out / f"recu-{slug}.docx"
            write_docx(receipt, path)
            written.append(path)

        files = "  ".join(str(p) for p in written)
        print(f"{receipt.number}  {money(receipt.total):>10} {CURRENCY}  {files}")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
